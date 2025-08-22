# pip install mcp pymupdf python-docx docx2pdf
# docx2pdf 는 Word/PowerPoint 설치 시에만 실제 변환 동작(없으면 자동 건너뜀)
# -*- coding: utf-8 -*-
"""
초간단 MCP Agent (Claude Desktop 데모용)
- 툴 2개:
  1) convert_to_pdf(input_path): .docx/.pptx -> PDF (CONVERTED/에 저장, 가능할 때만)
  2) search_docs(query, root_dir?): .pdf/.docx/.txt에서 키워드 스니펫 검색

설계 포인트:
- 의존성 최소화, 예외는 "부드럽게" 처리(데모 중 에러로 멈추지 않게)
- 변환 실패해도 검색은 원문에서 진행 가능
"""

import re
from pathlib import Path
from typing import Optional, List, Dict, Any

from mcp.server.fastmcp import FastMCP
mcp = FastMCP("doc-demo")

# ------------------------
# 유틸
# ------------------------
def downloads_dir() -> Path:
    return Path.home() / "Downloads"

def converted_dir(p: Path) -> Path:
    out = p.parent / "CONVERTED"
    out.mkdir(parents=True, exist_ok=True)
    return out

def read_txt(p: Path) -> str:
    for enc in ("utf-8", "cp949", "euc-kr"):
        try:
            return p.read_text(encoding=enc, errors="ignore")
        except Exception:
            pass
    try:
        return p.read_bytes().decode("utf-8", errors="ignore")
    except Exception:
        return ""

# ------------------------
# 변환(.docx/.pptx -> .pdf)
# ------------------------
def try_convert_to_pdf(src: Path) -> Optional[Path]:
    if src.suffix.lower() not in {".docx", ".pptx"}:
        return None
    try:
        from docx2pdf import convert  # Word/PowerPoint 설치 필요
    except Exception:
        return None

    out_pdf = converted_dir(src) / (src.stem + ".pdf")
    if out_pdf.exists():
        return out_pdf
    try:
        # 파일 단위 변환 후 같은 폴더에 생긴 PDF를 CONVERTED/로 이동
        convert(str(src))
        created = src.with_suffix(".pdf")
        if created.exists():
            created.replace(out_pdf)
            return out_pdf
    except Exception:
        pass
    return None

# ------------------------
# 텍스트 추출
# ------------------------
def extract_pdf(p: Path) -> str:
    try:
        import fitz  # PyMuPDF
        text_all = []
        with fitz.open(str(p)) as doc:
            for page in doc:
                text_all.append(page.get_text("text") or "")
        return "\n".join(text_all)
    except Exception:
        return ""  # PDF 미지원 환경이면 빈 문자열

def extract_docx(p: Path) -> str:
    try:
        import docx
        d = docx.Document(str(p))
        parts = [para.text for para in d.paragraphs]
        # 표(간단)
        for t in d.tables:
            for r in t.rows:
                parts.append("\t".join(c.text for c in r.cells))
        return "\n".join(parts)
    except Exception:
        return ""

def extract_text(p: Path) -> str:
    ext = p.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(p)
    if ext == ".docx":
        return extract_docx(p)
    if ext == ".txt":
        return read_txt(p)
    return ""  # 그 외 확장자는 이번 데모에서 제외

def make_snippets(text: str, needle: str, case_sensitive: bool, context=60, max_hits=5) -> List[str]:
    if not text:
        return []
    hay = text if case_sensitive else text.lower()
    ndl = needle if case_sensitive else needle.lower()
    out, idx = [], 0
    while True:
        idx = hay.find(ndl, idx)
        if idx == -1:
            break
        start, end = max(0, idx - context), min(len(text), idx + len(needle) + context)
        out.append(text[start:end].replace("\n", " "))
        if len(out) >= max_hits:
            break
        idx += len(needle)
    return out

# ------------------------
# MCP TOOLS
# ------------------------
@mcp.tool()
def convert_to_pdf(input_path: str) -> Dict[str, Any]:
    """
    .docx/.pptx -> PDF (CONVERTED/에 저장)
    - Word/PowerPoint가 없으면 조용히 실패(ok=False, 메시지 안내)
    """
    p = Path(input_path).expanduser().resolve()
    if not p.exists():
        return {"ok": False, "message": "입력 파일이 없습니다.", "output": None}
    if p.suffix.lower() not in {".docx", ".pptx"}:
        return {"ok": False, "message": "지원 확장자 아님(.docx/.pptx)", "output": None}

    out_pdf = try_convert_to_pdf(p)
    if out_pdf and out_pdf.exists():
        return {"ok": True, "message": "변환 완료", "output": str(out_pdf)}
    return {"ok": False, "message": "변환 불가(Word/PowerPoint 미설치 또는 docx2pdf 미사용 가능)", "output": None}

@mcp.tool()
def search_docs(
    query: str,
    root_dir: Optional[str] = None,
    case_sensitive: bool = False,
    max_results: int = 20
) -> Dict[str, Any]:
    """
    간단 검색: root_dir(기본 Downloads)에서 .pdf/.docx/.txt 파일을 순회하여 스니펫 반환
    - 변환이 필요하면 미리 convert_to_pdf로 만들어 CONVERTED/에 넣어두세요.
    - 데모 단순화를 위해 regex, 페이지 번호 등은 생략
    """
    base = Path(root_dir).expanduser().resolve() if root_dir else downloads_dir()
    if not base.exists():
        return {"ok": False, "message": f"루트 폴더 없음: {base}", "results": []}

    exts = {".pdf", ".docx", ".txt"}
    results = []
    scanned = 0

    for p in base.rglob("*"):
        if not p.is_file() or p.suffix.lower() not in exts:
            continue
        scanned += 1
        text = extract_text(p)
        snippets = make_snippets(text, query, case_sensitive)
        if snippets:
            results.append({"path": str(p), "snippets": snippets})
            if len(results) >= max_results:
                break

    return {
        "ok": True,
        "message": "검색 완료",
        "summary": {"root": str(base), "scanned": scanned, "hits": len(results)},
        "results": results
    }

# ------------------------
# ENTRY
# ------------------------
if __name__ == "__main__":
    mcp.run()
