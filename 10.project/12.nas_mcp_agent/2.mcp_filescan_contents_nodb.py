# mcp_filescan_no_db.py — 무DB 스캐너 (PDF/Docx 본문 미리보기/검색 지원)
# pip install mcp python-dotenv chardet pymupdf python-docx

import os, sys, fnmatch, json, re
import chardet
from pathlib import Path
from dotenv import load_dotenv
from mcp.server.fastmcp import FastMCP

# ===== PDF/Docx 추출기 준비 =====
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except Exception:
    PYMUPDF_AVAILABLE = False

try:
    import docx  # python-docx
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# ===== 콘솔 인코딩 =====
sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")

# ===== 환경 =====
# ENV_PATH가 지정돼 있으면 우선 로드
load_dotenv(dotenv_path=os.getenv("ENV_PATH"))
load_dotenv()

ROOT = Path(os.getenv("SCAN_ROOT", r"Z:\\")).resolve()  # 예: Z:\  (주의: 역슬래시는 2개)
mcp = FastMCP("NAS-NoDB-Scanner")

# ===== 파일 유형 =====
ALLOWED_EXT = {
    "pdf","docx","doc","xlsx","xls","pptx","ppt",
    "txt","md","csv","json","yaml","yml","html","htm","rtf",
    "odt","ods","odp"
}
# 텍스트 바이너리 아님(그대로 읽기 가능한 형)
PLAIN_TEXT_EXT = {"txt","md","csv","json","yaml","yml","html","htm","rtf"}

# ===== 공통 유틸 =====
def _is_within(root: Path, target: Path) -> bool:
    root = root.resolve()
    try:
        # Py3.9+: 안전
        return target.resolve().is_relative_to(root)
    except AttributeError:
        # Fallback
        try:
            return os.path.commonpath([str(root), str(target.resolve())]) == str(root)
        except ValueError:
            return False

def _ensure_within_root(p: Path):
    if not _is_within(ROOT, p):
        return {
            "error": "permission_denied",
            "message": f"SCAN_ROOT({ROOT}) 바깥은 접근할 수 없습니다.",
            "allowed_roots": [str(ROOT)]
        }
    if not p.exists():
        return {
            "error": "not_found",
            "message": f"경로가 존재하지 않습니다: {p}",
            "allowed_roots": [str(ROOT)]
        }
    return None

def _iter_files(base: Path, max_items: int = 10000, ext_filter=None, name_glob=None):
    n = 0
    for p in base.rglob("*"):
        try:
            if not p.is_file():
                continue
            ext = p.suffix.lower().lstrip(".")
            if ext_filter and ext not in ext_filter:
                continue
            if name_glob and not fnmatch.fnmatch(p.name, name_glob):
                continue
            stat = p.stat()
            yield {
                "path": str(p),
                "name": p.name,
                "ext": ext,
                "size": stat.st_size,
                "mtime": stat.st_mtime,
            }
            n += 1
            if n >= max_items:
                break
        except Exception:
            continue

def _detect_text_bytes_to_str(data: bytes) -> str:
    enc = (chardet.detect(data).get("encoding") or "utf-8").lower()
    return data.decode(enc, errors="ignore")

# ===== 본문 추출기 =====
def extract_text_plain(path: Path, max_bytes: int = 1_000_000) -> str:
    data = path.read_bytes()[:max_bytes]
    return _detect_text_bytes_to_str(data)

def extract_text_pdf(path: Path, max_pages: int = 10, max_chars: int = 200_000) -> str:
    if not PYMUPDF_AVAILABLE:
        return "(PyMuPDF 미설치: pip install pymupdf)"
    text_parts = []
    try:
        with fitz.open(str(path)) as doc:
            pages = min(len(doc), max_pages)
            for i in range(pages):
                page = doc.load_page(i)
                text_parts.append(page.get_text("text"))
                if sum(len(t) for t in text_parts) >= max_chars:
                    break
    except Exception as e:
        return f"(PDF 텍스트 추출 실패: {e})"
    return "".join(text_parts)[:max_chars]

def extract_text_docx(path: Path, max_chars: int = 200_000) -> str:
    if not DOCX_AVAILABLE:
        return "(python-docx 미설치: pip install python-docx)"
    try:
        d = docx.Document(str(path))
        parts = []
        # 본문
        parts.extend(p.text for p in d.paragraphs if p.text)
        # 표 셀 텍스트
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        text = "\n".join(parts)
        return text[:max_chars]
    except Exception as e:
        return f"(DOCX 텍스트 추출 실패: {e})"

def extract_text_any(path: Path,
                     max_bytes_plain: int = 1_000_000,
                     pdf_max_pages: int = 10,
                     max_chars: int = 200_000) -> str:
    ext = path.suffix.lower().lstrip(".")
    if ext in PLAIN_TEXT_EXT:
        return extract_text_plain(path, max_bytes=max_bytes_plain)
    if ext == "pdf":
        return extract_text_pdf(path, max_pages=pdf_max_pages, max_chars=max_chars)
    if ext == "docx":
        return extract_text_docx(path, max_chars=max_chars)
    # 그 외는 본문 미지원
    return "(미지원 형식)"

# ===== MCP 도구 =====
@mcp.tool()
def scan(dir: str = "",
         limit: int = 2000,
         types: str = "pdf,docx,md,txt",
         name_glob: str = "") -> str:
    """
    SCAN_ROOT 기준으로 풀스캔(무DB)하여 파일 목록을 나열합니다.
    - dir: ROOT 하위 경로(예: '팀문서/보고서')
    - limit: 최대 개수 (기본 2000)
    - types: 쉼표구분 확장자 필터
    - name_glob: 와일드카드(예: '*보고서*.pdf')
    """
    base = (ROOT / dir)
    err = _ensure_within_root(base)
    if err: return json.dumps(err, ensure_ascii=False)
    
    ext_filter = {t.strip().lstrip(".").lower() for t in types.split(",") if t.strip()} or None
    rows = list(_iter_files(base, max_items=int(limit), ext_filter=ext_filter, name_glob=name_glob or None))
    return json.dumps(rows, ensure_ascii=False)

@mcp.tool()
def head(path: str,
         n: int = 40,
         pdf_max_pages: int = 3,
         max_chars: int = 100_000,
         max_bytes_plain: int = 512_000) -> str:
    """
    파일 앞부분 미리보기(텍스트/HTML/RTF는 그대로, PDF/Docx는 추출 후).
    - n: 보여줄 줄 수
    - pdf_max_pages: PDF 추출 최대 페이지
    """
    p = Path(path)
    err = _ensure_within_root(p)
    if err: return json.dumps(err, ensure_ascii=False)
    
    try:
        text = extract_text_any(p, max_bytes_plain=max_bytes_plain,
                                pdf_max_pages=pdf_max_pages, max_chars=max_chars)
        lines = text.splitlines()[:max(1, n)]
        return "\n".join(lines) if lines else "(내용 없음)"
    except Exception as e:
        return f"(미리보기 실패: {e})"

@mcp.tool()
def grep(q: str,
         dir: str = "",
         types: str = "pdf,docx,txt,md,csv,json,yaml,yml,html,htm,rtf",
         limit: int = 200,            # 결과(라인) 수
         name_glob: str = "",
         pdf_max_pages: int = 10,
         max_chars: int = 200_000,
         max_bytes_plain: int = 1_000_000,
         regex: bool = False,
         ignore_case: bool = True) -> str:
    """
    무DB 즉시 본문 검색. 텍스트 + PDF + Docx 지원(느릴 수 있음).
    - q: 검색어(또는 정규식)
    - dir: SCAN_ROOT 하위 경로
    - types: 대상 확장자 필터
    - name_glob: 파일명 와일드카드
    - limit: 반환할 (파일,줄) 결과 수
    - pdf_max_pages/max_chars/max_bytes_plain: 추출 한도(성능 보호)
    - regex/ignore_case: 검색 방식 옵션
    """
    base = (ROOT / dir)
    err = _ensure_within_root(base)
    if err: return json.dumps(err, ensure_ascii=False)
    
    type_set = {t.strip().lstrip(".").lower() for t in types.split(",") if t.strip()}
    results = []

    # 패턴 생성
    if regex:
        flags = re.IGNORECASE if ignore_case else 0
        pat = re.compile(q, flags)
        def match_line(line): return pat.search(line) is not None
    else:
        needle = q.casefold() if ignore_case else q
        def match_line(line):
            hay = line.casefold() if ignore_case else line
            return needle in hay

    for info in _iter_files(base, max_items=1_000_000, ext_filter=type_set, name_glob=name_glob or None):
        p = Path(info["path"])
        try:
            text = extract_text_any(p, max_bytes_plain=max_bytes_plain,
                                    pdf_max_pages=pdf_max_pages, max_chars=max_chars)
            # 추출 불가/미지원 메시지인 경우 건너뜀
            if text.startswith("(") and "미설치" in text or "미지원" in text or "실패" in text:
                continue
            for i, line in enumerate(text.splitlines(), 1):
                if match_line(line):
                    snippet = line.strip()
                    if len(snippet) > 240:
                        snippet = snippet[:240] + "…"
                    results.append({"path": str(p), "line": i, "snippet": snippet})
                    if len(results) >= limit:
                        return json.dumps(results, ensure_ascii=False)
        except Exception:
            continue

    return json.dumps(results, ensure_ascii=False)

@mcp.tool()
def allowed_roots() -> str:
    return json.dumps({"allowed_roots":[str(ROOT)]}, ensure_ascii=False)

if __name__ == "__main__":
    print(f"[start] ROOT={ROOT}")
    mcp.run()
