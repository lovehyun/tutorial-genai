# mcp_filescan_contents_nodb.py
# - 무DB 즉시 스캔 도구(scan/head/grep) + 수동 인덱싱(index_new) + DB 검색(search_db)
# 의존성:
#   pip install mcp python-dotenv chardet pymupdf python-docx
# (PDF/Docx 추출 비사용이면 pymupdf/python-docx 생략 가능)

import os, sys, fnmatch, json, re, mimetypes, sqlite3, time
import chardet
from pathlib import Path
from dotenv import load_dotenv
from mcp.server.fastmcp import FastMCP

# ===== PDF/Docx 추출기 =====
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except Exception:
    PYMUPDF_AVAILABLE = False

try:
    import docx  # python-docx
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# ===== 콘솔 인코딩 =====
sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")

# ===== 환경 =====
load_dotenv(dotenv_path=os.getenv("ENV_PATH"))
load_dotenv()

ROOT = Path(os.getenv("SCAN_ROOT", r"Z:\\")).resolve()  # 기본 Z:\  (역슬래시는 2개)
INDEX_DB = Path(os.getenv("INDEX_DB", "nas_index.db")).resolve()  # DB 경로(기본: 현재 폴더)
mcp = FastMCP("NAS-NoDB-Scanner+Index")

ALLOWED_EXT = {
    "pdf","docx","doc","xlsx","xls","pptx","ppt",
    "txt","md","csv","json","yaml","yml","html","htm","rtf",
    "odt","ods","odp"
}
PLAIN_TEXT_EXT = {"txt","md","csv","json","yaml","yml","html","htm","rtf"}

# ===== 공통 유틸 =====
def _is_within(root: Path, target: Path) -> bool:
    root = root.resolve()
    try:
        # Py3.9+: 안전
        return target.resolve().is_relative_to(root)
    except AttributeError:
        # Fallback
        try:
            return os.path.commonpath([str(root), str(target.resolve())]) == str(root)
        except ValueError:
            return False
        
def _ensure_within_root(p: Path):
    if not _is_within(ROOT, p):
        return {
            "error": "permission_denied",
            "message": f"SCAN_ROOT({ROOT}) 바깥은 접근할 수 없습니다.",
            "allowed_roots": [str(ROOT)]
        }
    if not p.exists():
        return {
            "error": "not_found",
            "message": f"경로가 존재하지 않습니다: {p}",
            "allowed_roots": [str(ROOT)]
        }
    return None

def _iter_files(base: Path, max_items: int = 10000, ext_filter=None, name_glob=None):
    n = 0
    for p in base.rglob("*"):
        try:
            if not p.is_file():
                continue
            ext = p.suffix.lower().lstrip(".")
            if ext_filter and ext not in ext_filter:
                continue
            if name_glob and not fnmatch.fnmatch(p.name, name_glob):
                continue
            st = p.stat()
            yield {
                "path": str(p),
                "name": p.name,
                "ext": ext,
                "size": st.st_size,
                "mtime": st.st_mtime,
                "mime": mimetypes.guess_type(str(p))[0] or "application/octet-stream",
            }
            n += 1
            if n >= max_items:
                break
        except Exception:
            continue

def _detect_text_bytes_to_str(data: bytes) -> str:
    enc = (chardet.detect(data).get("encoding") or "utf-8").lower()
    return data.decode(enc, errors="ignore")

# ===== 본문 추출 =====
def extract_text_plain(path: Path, max_bytes: int = 1_000_000) -> str:
    data = path.read_bytes()[:max_bytes]
    return _detect_text_bytes_to_str(data)

def extract_text_pdf(path: Path, max_pages: int = 10, max_chars: int = 200_000) -> str:
    if not PYMUPDF_AVAILABLE:
        return "(PyMuPDF 미설치: pip install pymupdf)"
    parts = []
    try:
        with fitz.open(str(path)) as doc:
            pages = min(len(doc), max_pages)
            for i in range(pages):
                page = doc.load_page(i)
                parts.append(page.get_text("text"))
                if sum(len(t) for t in parts) >= max_chars:
                    break
    except Exception as e:
        return f"(PDF 텍스트 추출 실패: {e})"
    return "".join(parts)[:max_chars]

def extract_text_docx(path: Path, max_chars: int = 200_000) -> str:
    if not DOCX_AVAILABLE:
        return "(python-docx 미설치: pip install python-docx)"
    try:
        d = docx.Document(str(path))
        parts = []
        parts.extend(p.text for p in d.paragraphs if p.text)
        for tbl in d.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts)[:max_chars]
    except Exception as e:
        return f"(DOCX 텍스트 추출 실패: {e})"

def extract_text_any(path: Path,
                     max_bytes_plain: int = 1_000_000,
                     pdf_max_pages: int = 10,
                     max_chars: int = 200_000) -> str:
    ext = path.suffix.lower().lstrip(".")
    if ext in PLAIN_TEXT_EXT:
        return extract_text_plain(path, max_bytes=max_bytes_plain)
    if ext == "pdf":
        return extract_text_pdf(path, max_pages=pdf_max_pages, max_chars=max_chars)
    if ext == "docx":
        return extract_text_docx(path, max_chars=max_chars)
    return "(미지원 형식)"

# ===== 무DB 도구: scan/head/grep =====
@mcp.tool()
def scan(dir: str = "",
         limit: int = 2000,
         types: str = "pdf,docx,md,txt",
         name_glob: str = "") -> str:
    base = (ROOT / dir)
    err = _ensure_within_root(base)
    if err: return json.dumps(err, ensure_ascii=False)
    ext_filter = {t.strip().lstrip(".").lower() for t in types.split(",") if t.strip()} or None
    rows = list(_iter_files(base, max_items=int(limit), ext_filter=ext_filter, name_glob=name_glob or None))
    return json.dumps(rows, ensure_ascii=False)

@mcp.tool()
def head(path: str,
         n: int = 40,
         pdf_max_pages: int = 3,
         max_chars: int = 100_000,
         max_bytes_plain: int = 512_000) -> str:
    p = Path(path)
    err = _ensure_within_root(p)
    if err: return json.dumps(err, ensure_ascii=False)
    
    try:
        text = extract_text_any(p, max_bytes_plain=max_bytes_plain,
                                pdf_max_pages=pdf_max_pages, max_chars=max_chars)
        lines = text.splitlines()[:max(1, n)]
        return "\n".join(lines) if lines else "(내용 없음)"
    except Exception as e:
        return f"(미리보기 실패: {e})"

@mcp.tool()
def grep(q: str,
         dir: str = "",
         types: str = "pdf,docx,txt,md,csv,json,yaml,yml,html,htm,rtf",
         limit: int = 200,
         name_glob: str = "",
         pdf_max_pages: int = 10,
         max_chars: int = 200_000,
         max_bytes_plain: int = 1_000_000,
         regex: bool = False,
         ignore_case: bool = True) -> str:
    base = (ROOT / dir)
    err = _ensure_within_root(base)
    if err: return json.dumps(err, ensure_ascii=False)
    
    type_set = {t.strip().lstrip(".").lower() for t in types.split(",") if t.strip()}
    results = []

    if regex:
        flags = re.IGNORECASE if ignore_case else 0
        pat = re.compile(q, flags)
        def match_line(line): return pat.search(line) is not None
    else:
        needle = q.casefold() if ignore_case else q
        def match_line(line):
            hay = line.casefold() if ignore_case else line
            return needle in hay

    for info in _iter_files(base, max_items=1_000_000, ext_filter=type_set, name_glob=name_glob or None):
        p = Path(info["path"])
        try:
            text = extract_text_any(p, max_bytes_plain=max_bytes_plain,
                                    pdf_max_pages=pdf_max_pages, max_chars=max_chars)
            if text.startswith("(") and ("미설치" in text or "미지원" in text or "실패" in text):
                continue
            for i, line in enumerate(text.splitlines(), 1):
                if match_line(line):
                    snippet = line.strip()
                    if len(snippet) > 240:
                        snippet = snippet[:240] + "…"
                    results.append({"path": str(p), "line": i, "snippet": snippet})
                    if len(results) >= limit:
                        return json.dumps(results, ensure_ascii=False)
        except Exception:
            continue
    return json.dumps(results, ensure_ascii=False)

@mcp.tool()
def allowed_roots() -> str:
    return json.dumps({"allowed_roots":[str(ROOT)]}, ensure_ascii=False)

@mcp.tool()
def where_db() -> str:
    return str(INDEX_DB)

# ===== DB(FTS5) — 수동 인덱싱 & 검색 =====
SCHEMA = """
PRAGMA journal_mode=WAL;
PRAGMA synchronous=NORMAL;
PRAGMA temp_store=MEMORY;
CREATE TABLE IF NOT EXISTS files(
  id INTEGER PRIMARY KEY,
  path TEXT UNIQUE NOT NULL,
  name TEXT NOT NULL,
  ext  TEXT,
  size INTEGER,
  mtime REAL,
  mime TEXT,
  indexed_at REAL
);
CREATE VIRTUAL TABLE IF NOT EXISTS files_fts USING fts5(
  name, path, content, content_rowid='id'
);
"""

class IndexDB:
    def __init__(self, db_path: Path):
        self.conn = sqlite3.connect(str(db_path))
        self.conn.executescript(SCHEMA)
        self.conn.commit()

    def upsert_meta_and_fts(self, meta: dict, content: str):
        cur = self.conn.cursor()
        cur.execute(
            "INSERT INTO files(path,name,ext,size,mtime,mime,indexed_at) VALUES(?,?,?,?,?,?,?) "
            "ON CONFLICT(path) DO UPDATE SET name=excluded.name, ext=excluded.ext, size=excluded.size, "
            "mtime=excluded.mtime, mime=excluded.mime, indexed_at=excluded.indexed_at",
            (meta["path"], meta["name"], meta["ext"], meta["size"], meta["mtime"], meta["mime"], time.time())
        )
        # rowid 확인
        cur.execute("SELECT id FROM files WHERE path=?", (meta["path"],))
        rowid = cur.fetchone()[0]
        # FTS5는 INSERT OR REPLACE 사용
        cur.execute(
            "INSERT OR REPLACE INTO files_fts(rowid,name,path,content) VALUES(?,?,?,?)",
            (rowid, meta["name"], meta["path"], content or "")
        )

    def exists_by_path(self, path: str) -> bool:
        cur = self.conn.cursor()
        cur.execute("SELECT 1 FROM files WHERE path=? LIMIT 1", (path,))
        return cur.fetchone() is not None

    def stat(self):
        cur = self.conn.cursor()
        cur.execute("SELECT COUNT(*), IFNULL(SUM(size),0) FROM files")
        c, s = cur.fetchone()
        cur.execute("SELECT ext, COUNT(*) FROM files GROUP BY ext ORDER BY COUNT(*) DESC")
        by_ext = cur.fetchall()
        return {"count": c or 0, "total_bytes": s or 0, "by_ext": by_ext}

    def reset(self):
        cur = self.conn.cursor()
        cur.executescript("DROP TABLE IF EXISTS files; DROP TABLE IF EXISTS files_fts;")
        self.conn.executescript(SCHEMA)
        self.conn.commit()

@mcp.tool()
def index_reset() -> str:
    """인덱스 DB를 초기화합니다."""
    db = IndexDB(INDEX_DB)
    db.reset()
    return "OK"

@mcp.tool()
def index_stat() -> str:
    """인덱스 통계를 확인합니다."""
    db = IndexDB(INDEX_DB)
    return json.dumps(db.stat(), ensure_ascii=False)

@mcp.tool()
def index_new(dir: str = "",
              types: str = "pdf,docx,md,txt",
              include: str = "none",   # 'none' | 'light'
              limit: int = 1_000_000,
              batch_commit: int = 200,
              pdf_max_pages: int = 10,
              max_chars: int = 200_000,
              max_bytes_plain: int = 1_000_000) -> str:
    """
    수동 인덱싱: DB에 없는 파일만 신규로 인덱스(이미 인덱싱된 파일은 건너뜀).
    - include='none': 파일명/경로만 FTS (가볍고 빠름)
    - include='light': 텍스트형+PDF/Docx 추출 본문(제한 크기)까지 FTS
    """
    base = (ROOT / dir).resolve()
    if not str(base).startswith(str(ROOT)) or not base.exists():
        return json.dumps({"error":"경로 없음 또는 ROOT 바깥 접근 불가"}, ensure_ascii=False)

    ext_filter = {t.strip().lstrip(".").lower() for t in types.split(",") if t.strip()} or None
    db = IndexDB(INDEX_DB)

    added = 0
    scanned = 0
    for meta in _iter_files(base, max_items=int(limit), ext_filter=ext_filter, name_glob=None):
        scanned += 1
        if db.exists_by_path(meta["path"]):
            continue  # 이미 인덱스된 파일은 건너뜀
        content = ""
        if include.lower() == "light":
            p = Path(meta["path"])
            text = extract_text_any(p, max_bytes_plain=max_bytes_plain,
                                    pdf_max_pages=pdf_max_pages, max_chars=max_chars)
            # 추출 실패/미설치 메시지는 비우고 넘어감(이름/경로만 인덱스)
            if not (text.startswith("(") and ("미설치" in text or "미지원" in text or "실패" in text)):
                content = text
        db.upsert_meta_and_fts(meta, content)
        added += 1
        if added % batch_commit == 0:
            db.conn.commit()

    db.conn.commit()
    return json.dumps({"scanned": scanned, "added": added, "include": include, "db": str(INDEX_DB)}, ensure_ascii=False)

@mcp.tool()
def search_db(q: str = "",
              types: str = "pdf,docx,md,txt",
              scope: str = "all",  # 'all' | 'name' | 'content'
              limit: int = 50) -> str:
    """
    DB(FTS5) 검색. q 비우면 최신 파일명/경로 순(ctime/mtime가 아닌 files.mtime 기준)으로 보여줌.
    - scope='name' → name/path 대상만
    - scope='content' → content 대상만(본문 인덱싱한 경우에 한함)
    """
    db = IndexDB(INDEX_DB)
    cur = db.conn.cursor()
    type_set = {t.strip().lstrip(".").lower() for t in types.split(",") if t.strip()}

    if not q.strip():
        # q 없으면 메타에서 최신순
        if type_set:
            placeholders = ",".join(["?"]*len(type_set))
            cur.execute(f"""
              SELECT f.path, f.name, f.ext, f.size, f.mtime, f.mime
              FROM files f
              WHERE f.ext IN ({placeholders})
              ORDER BY f.mtime DESC
              LIMIT ?
            """, (*type_set, int(limit)))
        else:
            cur.execute("""
              SELECT f.path, f.name, f.ext, f.size, f.mtime, f.mime
              FROM files f
              ORDER BY f.mtime DESC
              LIMIT ?
            """, (int(limit),))
        rows = cur.fetchall()
        return json.dumps([{
            "path": r[0], "name": r[1], "ext": r[2], "size": r[3], "mtime": r[4], "mime": r[5]
        } for r in rows], ensure_ascii=False)

    # FTS MATCH 구성
    # scope에 따라 필드 한정
    fts_q = q
    if scope == "name":
        fts_q = f'(name:{q}) OR (path:{q})'
    elif scope == "content":
        fts_q = f'(content:{q})'

    if type_set:
        placeholders = ",".join(["?"]*len(type_set))
        cur.execute(f"""
          SELECT f.path, f.name, f.ext, f.size, f.mtime, f.mime
          FROM files f
          JOIN files_fts ft ON f.id = ft.rowid
          WHERE ft MATCH ?
            AND f.ext IN ({placeholders})
          ORDER BY bm25(ft)
          LIMIT ?
        """, (fts_q, *type_set, int(limit)))
    else:
        cur.execute("""
          SELECT f.path, f.name, f.ext, f.size, f.mtime, f.mime
          FROM files f
          JOIN files_fts ft ON f.id = ft.rowid
          WHERE ft MATCH ?
          ORDER BY bm25(ft)
          LIMIT ?
        """, (fts_q, int(limit)))
    rows = cur.fetchall()
    return json.dumps([{
        "path": r[0], "name": r[1], "ext": r[2], "size": r[3], "mtime": r[4], "mime": r[5]
    } for r in rows], ensure_ascii=False)

if __name__ == "__main__":
    print(f"[start] ROOT={ROOT} | DB={INDEX_DB}")
    mcp.run()
